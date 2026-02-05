import sys
import os
from docx import Document

# Add backend directory to sys.path
current_dir = os.path.dirname(os.path.abspath(__file__))
backend_dir = os.path.dirname(current_dir)
sys.path.append(backend_dir)

def inspect_indentation_sales():
    file_path = r"D:\workspace\DocComparison\买卖合同(销售).docx"
    doc = Document(file_path)
    
    print(f"\n--- Inspecting {os.path.basename(file_path)} '四、验收' ---")
    
    found_section = False
    count = 0
    for p in doc.paragraphs:
        txt = p.text.strip()
        if "四、" in txt and "验收" in txt:
            found_section = True
            print(f"HEADER: {txt}")
            continue
        
        if found_section:
            if "五、" in txt: # Next section
                break
                
            if not txt: continue
            
            # Print indentation details
            pf = p.paragraph_format
            left = pf.left_indent.pt if pf.left_indent else 0
            first = pf.first_line_indent.pt if pf.first_line_indent else 0
            
            # Check numbering
            num_id = "None"
            ilvl = "None"
            if p._element.pPr is not None and p._element.pPr.numPr is not None:
                if p._element.pPr.numPr.numId is not None:
                    num_id = p._element.pPr.numPr.numId.val
                if p._element.pPr.numPr.ilvl is not None:
                    ilvl = p._element.pPr.numPr.ilvl.val
            
            print(f"Item: {txt[:20]}...")
            print(f"  -> Left: {left} pt, First: {first} pt")
            print(f"  -> NumId: {num_id}, ilvl: {ilvl}")
            
            count += 1
            if count > 5: break

def inspect_indentation_confidentiality():
    file_path = r"D:\workspace\DocComparison\保密协议_双方-范本.docx"
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
        
    doc = Document(file_path)
    
    print(f"\n--- Inspecting {os.path.basename(file_path)} '定义' ---")
    
    count = 0
    start_tracking = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        
        if "定义：" in txt and not start_tracking:
            start_tracking = True
            
        if start_tracking:
            pf = p.paragraph_format
            left = pf.left_indent.pt if pf.left_indent else 0
            first = pf.first_line_indent.pt if pf.first_line_indent else 0
            
            print(f"Text: {txt[:20]}...")
            print(f"  -> Left: {left} pt, First: {first} pt")
            
            count += 1
            if count > 10: break

if __name__ == "__main__":
    inspect_indentation_sales()
    inspect_indentation_confidentiality()
