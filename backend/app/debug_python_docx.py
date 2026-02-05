import docx
from docx.shared import Pt
import sys
import os

sys.path.append('/app')

def analyze_docx_structure():
    file_path = "/app/app/test_doc_sales.docx"
    output_path = "/app/app/python_docx_analysis_sales.txt"
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    print(f"Analyzing {file_path} with python-docx...")
    
    try:
        doc = docx.Document(file_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Total Paragraphs: {len(doc.paragraphs)}\n")
            f.write("="*50 + "\n")
            
            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if not text:
                    continue
                    
                style = p.style.name
                
                # Numbering properties
                num_id = None
                ilvl = None
                if p._element.pPr is not None and p._element.pPr.numPr is not None:
                    numPr = p._element.pPr.numPr
                    if numPr.numId is not None:
                        num_id = numPr.numId.val
                    if numPr.ilvl is not None:
                        ilvl = numPr.ilvl.val
                
                # Indentation
                indent_left = p.paragraph_format.left_indent
                indent_first = p.paragraph_format.first_line_indent
                
                indent_info = []
                if indent_left: indent_info.append(f"Left: {indent_left.pt}pt")
                if indent_first: indent_info.append(f"First: {indent_first.pt}pt")
                
                f.write(f"Para {i}:\n")
                f.write(f"  Text: {text[:100]!r}\n") # Truncate long text
                f.write(f"  Style: {style}\n")
                if num_id is not None or ilvl is not None:
                    f.write(f"  Numbering: ID={num_id}, Level={ilvl}\n")
                if indent_info:
                    f.write(f"  Indent: {', '.join(indent_info)}\n")
                
                # Check for run-level formatting that might imply headings (bold, font size)
                runs_info = []
                for run in p.runs:
                    props = []
                    if run.bold: props.append("Bold")
                    if run.font.size: props.append(f"Size:{run.font.size.pt}")
                    if run.underline: props.append(f"Underline:{run.underline}")
                    if run.font.underline: props.append(f"FontUnderline:{run.font.underline}")
                    if props:
                        runs_info.append(f"{run.text[:30]!r}:{','.join(props)}")
                
                if runs_info:
                    f.write(f"  Runs: {' | '.join(runs_info)}\n")
                
                f.write("-" * 30 + "\n")

            try:
                from app.services.doc_service import DocService
                from app.services.diff_service import align_blocks
                import copy
                blocks = DocService.parse_docx(file_path)
                b = next((x for x in blocks if (x.text and "运输方式" in x.text)), None)
                f.write("\n" + "="*50 + "\n")
                f.write("DocService HTML Preview\n")
                f.write("="*50 + "\n")
                f.write(f"Found: {bool(b)}\n")
                if b is not None:
                    f.write(f"BlockText: {b.text!r}\n")
                    f.write(f"BlockHtml: {b.htmlFragment}\n")

                    right_blocks = copy.deepcopy(blocks)
                    rb = next((x for x in right_blocks if x.blockId == b.blockId), None)
                    if rb is not None:
                        lines = (rb.text or "").split("\n")
                        if lines:
                            lines[0] = lines[0] + "X"
                            rb.text = "\n".join(lines)
                        rows = align_blocks(blocks, right_blocks)
                        changed = next((r for r in rows if r.kind == "changed" and r.leftBlockId == b.blockId), None)
                        f.write("\nChanged Diff Preview\n")
                        f.write(f"ChangedFound: {bool(changed)}\n")
                        if changed is not None:
                            f.write(f"HasUnderlineSpan: {'text-decoration: underline' in (changed.rightDiffHtml or '')}\n")
                            f.write(f"HasTransportText: {'运输方式' in (changed.rightDiffHtml or '')}\n")
            except Exception as e:
                f.write("\nDocService parse failed: " + repr(e) + "\n")

            f.write("\n" + "="*50 + "\n")
            f.write(f"Total Tables: {len(doc.tables)}\n")
            f.write("="*50 + "\n")
            
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i}: {len(table.rows)} rows x {len(table.columns)} cols\n")
                for r_idx, row in enumerate(table.rows):
                    row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    f.write(f"  Row {r_idx}: {row_cells}\n")
                f.write("-" * 30 + "\n")
                
        print(f"Analysis complete. Output written to {output_path}")
        
    except Exception as e:
        print(f"Error analyzing docx: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_docx_structure()
