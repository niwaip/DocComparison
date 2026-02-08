import os
import tempfile
import unittest


class DocParsingPurchaseContractTests(unittest.TestCase):
    def test_purchase_contract_contains_key_sections(self):
        from app.services.doc_service import DocService
        from docx import Document

        repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        candidates = [
            os.path.join(repo_root, "standard-contracts", "purchase.docx"),
            os.path.join(repo_root, "standard-contracts", "sales.docx"),
            os.path.join(repo_root, "买卖合同(销售).docx"),
            os.path.join(repo_root, "买卖合同(采购).docx"),
        ]

        def extract_raw_text(path: str) -> str:
            doc = Document(path)
            parts: list[str] = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = (cell.text or "").strip()
                        if t:
                            parts.append(t.replace("\n", " "))
            return "\n".join(parts)

        targets = ["验收", "保密", "知识产权", "贿赂", "反腐败", "出口", "审计"]
        for doc_path in candidates:
            if not os.path.exists(doc_path):
                continue
            with self.subTest(doc=os.path.basename(doc_path)):
                raw_text = extract_raw_text(doc_path)
                blocks = DocService.parse_docx(doc_path)
                parsed_text = "\n".join((b.text or "") for b in blocks)

                self.assertTrue(len(parsed_text.strip()) > 0)

                for kw in targets:
                    if kw in raw_text:
                        self.assertIn(kw, parsed_text)

    def test_purchase_contract_keeps_delivery_section_in_one_block(self):
        from app.services.doc_service import DocService
        from docx import Document

        repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        doc_path = os.path.join(repo_root, "买卖合同(采购).docx")
        if not os.path.exists(doc_path):
            self.skipTest("买卖合同(采购).docx not found")

        raw_text = "\n".join((p.text or "").strip() for p in Document(doc_path).paragraphs if (p.text or "").strip())
        blocks = DocService.parse_docx(doc_path)

        if "交货方式" not in raw_text:
            self.skipTest("delivery section not found in document")

        must_in_same_block = ["交货方式", "运输方式", "交货地址", "交货日期", "最终用户", "运费"]
        hit = None
        for b in blocks:
            bt = b.text or ""
            if all(k in bt for k in must_in_same_block):
                hit = b
                break
        self.assertIsNotNone(hit)

    def test_purchase_contract_keeps_appendix_section_in_one_block(self):
        from app.services.doc_service import DocService
        from docx import Document

        repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        doc_path = os.path.join(repo_root, "买卖合同(采购).docx")
        if not os.path.exists(doc_path):
            self.skipTest("买卖合同(采购).docx not found")

        raw_text = "\n".join((p.text or "").strip() for p in Document(doc_path).paragraphs if (p.text or "").strip())
        blocks = DocService.parse_docx(doc_path)

        if "合同附则" not in raw_text:
            self.skipTest("appendix section not found in document")

        must_in_same_block = ["合同附则", "税费", "附件一", "附件二"]
        hit = None
        for b in blocks:
            bt = b.text or ""
            if all(k in bt for k in must_in_same_block):
                hit = b
                break
        self.assertIsNotNone(hit)

    def test_purchase_contract_no_large_tail_duplication(self):
        from app.services.doc_service import DocService

        repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        doc_path = os.path.join(repo_root, "买卖合同(采购).docx")
        if not os.path.exists(doc_path):
            self.skipTest("买卖合同(采购).docx not found")

        blocks = DocService.parse_docx(doc_path)
        parsed_text = "\n".join((b.text or "") for b in blocks)

        markers = [
            "乙方应按本条上述记载的内容将本合同项下产品送至指定的交货地点及收货人",
            "本合同正本一式",
            "合同经甲乙双方盖章生效",
        ]
        for s in markers:
            if s in parsed_text:
                self.assertEqual(parsed_text.count(s), 1)

    def test_purchase_contract_html_has_no_indent_styles(self):
        from app.services.doc_service import DocService

        repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
        doc_path = os.path.join(repo_root, "买卖合同(采购).docx")
        if not os.path.exists(doc_path):
            self.skipTest("买卖合同(采购).docx not found")

        blocks = DocService.parse_docx(doc_path)
        html_all = "".join((b.htmlFragment or "") for b in blocks)
        self.assertNotIn("padding-left", html_all)
        self.assertNotIn("text-indent", html_all)
        self.assertNotIn("data-left-pt", html_all)
        self.assertNotIn("data-first-pt", html_all)

    def test_backend_does_not_inject_leading_spaces_for_indent(self):
        from app.services.doc_service import DocService
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, "t.docx")
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("  ")
            para.add_run("六、保密条款：")
            u = para.add_run("______")
            u.underline = True
            doc.save(p)

            blocks = DocService.parse_docx(p)
            html_all = "".join((b.htmlFragment or "") for b in blocks)
            self.assertIn("<p>六、保密条款：", html_all)
            self.assertNotIn("<p>  六、保密条款：", html_all)

    def test_footer_page_number_is_not_extracted_as_block(self):
        from app.services.doc_service import DocService
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            p = os.path.join(td, "t.docx")
            doc = Document()
            doc.add_paragraph("正文")
            sec = doc.sections[0]
            footer = sec.footer
            if footer.paragraphs:
                footer.paragraphs[0].text = "2"
            else:
                footer.add_paragraph("2")
            doc.save(p)

            blocks = DocService.parse_docx(p)
            texts = [(b.text or "").strip() for b in blocks]
            self.assertNotIn("2", texts)
            self.assertTrue(all("2" != line.strip() for t in texts for line in t.split("\n")))

    def test_diff_preserves_block_line_wrappers_for_indent(self):
        from app.models import Block, BlockKind, BlockMeta
        from app.services.diff_service import align_blocks

        left = Block(
            blockId="l1",
            kind=BlockKind.PARAGRAPH,
            structurePath="body.p[0]",
            stableKey="l1",
            text="8.2 保密期限为接收方收到相关专有信息起 年。\n9. 适用法律：",
            htmlFragment=(
                "<p>8.2 保密期限为接收方收到相关专有信息起 年。</p>"
                "<p>9. 适用法律：</p>"
            ),
            meta=BlockMeta(),
        )
        right = Block(
            blockId="r1",
            kind=BlockKind.PARAGRAPH,
            structurePath="body.p[0]",
            stableKey="r1",
            text="8.2 保密期限为接收方收到相关专有信息起 3 年。\n9. 适用法律：",
            htmlFragment=(
                "<p>8.2 保密期限为接收方收到相关专有信息起 3 年。</p>"
                "<p>9. 适用法律：</p>"
            ),
            meta=BlockMeta(),
        )

        rows = align_blocks([left], [right])
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0].kind, "changed")
        self.assertIsNotNone(rows[0].leftDiffHtml)
        self.assertIsNotNone(rows[0].rightDiffHtml)
        self.assertIn("aligned-table", rows[0].leftDiffHtml or "")
        self.assertIn("<p>", rows[0].leftDiffHtml or "")
        self.assertIn("<p>", rows[0].rightDiffHtml or "")
